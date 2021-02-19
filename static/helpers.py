import os
THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))

class Client :
    #Initialise a client with appropriate details
    def __init__(self, first_name, last_name, dva_num, address, suburb, state, postcode, mobile_number, home_number):
        self.first_name = first_name
        self.last_name = last_name
        self.dva_num = dva_num
        self.address = address
        self.suburb = suburb
        self.state = state
        self.postcode = postcode
        self.mobile_number = mobile_number
        self.home_number = home_number
    # Deliveries text
    def update_doc(self):
        from docx import Document
        document = Document(os.path.join(THIS_FOLDER, './static/deliveries.docx'))
        text = f"{self.first_name} {self.last_name} \n{self.address}, {self.suburb} \n"
        text += f"Mobile: {self.mobile_number}\n" if self.mobile_number else ""
        text += f"Home: {self.home_number}\n" if self.home_number else ""
        document.add_paragraph(text)
        document.save(os.path.join(THIS_FOLDER, './static/deliveries.docx'))

def make_doc():
    from docx import Document
    from datetime import date
    current_date = date.today()
    document = Document()
    document.add_heading("Deliveres for " + current_date.strftime("%d/%m/%Y"))
    document.save(os.path.join(THIS_FOLDER, './static/deliveries.docx'))

def get_text():
    from docx import Document
    document = Document(os.path.join(THIS_FOLDER, './static/deliveries.docx'))
    text = []
    for para in document.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

class Product:
    #Initialise a product with appropriate details
    def __init__(self, reference, lot, quantity, description):
        self.reference = reference
        self.lot = lot
        self.quantity = quantity
        self.description = description

class Form:
    #Initialise a form with a Client, pages of Products, and options
    def __init__(self, client, products, options, new, page_options):
        #Assign main details to values
        self.new = new
        self.client = client
        self.details = [
        client.dva_num,
        client.first_name[0],
        client.last_name,
        client.suburb,
        client.state,
        client.postcode,
        ""]

        #Product codes for frequent service products
        SERVICE_PRODUCTS = {
            "report": Product("REPORT-PAP", "SERVICE CLINICAL", "1", "PAP COMPLIANCE DOWNLOAD REPORT"),
            "visit": Product("VISIT-PAP", "SERVICE CLINICAL", "1", "PAP CONSULTATION"),
            "delivery": Product(self.find_distance(), "SERVICE TRAVEL", "1", "DELIVERY"),
            "setup": Product("SETUP-PAP", "SERVICE CLINICAL", "1", "PAP INITIAL SETUP AND 2 X FOLLOW UP"),
            "urgent": Product("URGENT-PAP", "SERVICE CLINICAL", "1", "URGENT DELIVERY")
        }

        #Format address appropriately across the three fields
        addresses = [client.address, "", ""]
        for (i, address) in enumerate(addresses):
            if len(address) > 25:
                for j in range(len(address.split())):
                    new = " ".join(address.split()[:-(j+1)])
                    if len(new) < 25 and not new == address:
                        addresses[i] = new
                        addresses[i+1] = " ".join(address.split()[-j-1:])
                        break
        self.details[3:3] = addresses

        pages = [[]]
        i = 0
        code_options = []
        for option, setting in options.items():
            if setting:
                code_options.append(SERVICE_PRODUCTS[option])
        if not products:
            pages[0] = code_options

        #Generate invoice text
        string = f"This invoice is for DVA:{client.dva_num}."
        for product in products:
            string += f" {product.quantity}x {product.description} REF: {product.reference} LOT:{product.lot},"
        self.text = string[:-1]

        #Populate "pages" with products, keeping services together on the same page
        while products:
            pages[i].append(products.pop(0))
            if not products:
                if (len(pages[i]) + len(code_options)) > 5:
                    pages.append(code_options)
                else:
                    pages[i] += code_options
                break
            if len(pages[i]) == 5:
                pages.append([])
                i += 1
        if page_options["phone-consult"]:
            pages.append([SERVICE_PRODUCTS["report"], SERVICE_PRODUCTS["visit"]])
        if page_options["phone-consult-vis"]:
            pages.append([SERVICE_PRODUCTS["visit"]])

        self.checklist = page_options["checklist"]

        #Assign values
        self.pages = pages

    def find_distance(self):
        #Import API libaries
        import googlemaps

        #Open file for secret
        secret = open(os.path.join(THIS_FOLDER, 'secret.txt'), "r").readlines()

        #Work address
        WORK_ADDRESS = secret[3]

        #Google maps API key
        GOOGLE_MAPS_API_KEY = secret[4]

        #Find distance between origin and form address using Google Maps API
        gmaps = googlemaps.Client(key=GOOGLE_MAPS_API_KEY)
        distances = gmaps.distance_matrix(WORK_ADDRESS, self.client.address + " " + self.client.suburb + " " + self.client.state)

        #Try to use distance to output appropriate reference
        try:
            dist = distances["rows"][0]["elements"][0]["distance"]["value"] * 2
            if dist < 50000:
                distance = 50
            elif dist >= 50000 and dist < 100000:
                distance = 100
            elif dist >= 100000 and dist < 200000:
                distance = 200
            elif dist >=200000:
                distance = 201

            return str(distance) + "DIST"

        #If address is not found, will return reference of "ERROR" and print an error message
        except:
            print("Location not found, as per matrix: \n", distances)
            return "ERROR"

    def make_pdf(self):
        #Import dependencies
        from PyPDF2 import PdfFileReader
        from datetime import date
        import os
        import pypdftk

        pdf_pages = []

        #Cycle through pages
        for j, page in enumerate(self.pages):

            template_name = os.path.join(THIS_FOLDER, "static/pdf_templates/form.pdf")

            #Read pdf templates using PyPDF2
            form = PdfFileReader(open(template_name, "rb"))

            #Get main form field names from pdf reader
            fields = form.getFields(tree=None, retval=None, fileobj=None)
            field_names = list(fields.keys())

            #Make a copy of field_values
            field_values = self.details[:]

            #Add values from each page
            for product in page:
                field_values += [product.reference, product.lot, product.quantity, product.description]

            #Pad out unused fields, zip into dict for writing
            field_values += [""] * (len(field_names) - len(field_values))
            field_dict = dict(zip(field_names, map(lambda x:x.upper(), field_values)))
            #Add page to writer, update fields from input data
            pdf_pages.append(pypdftk.fill_form(template_name, field_dict))

        if self.checklist:
            end_form_template_name = os.path.join(THIS_FOLDER, "static/pdf_templates/end_page.pdf")
            #Get pdf templates using PyPDF2
            end_form = PdfFileReader(open(end_form_template_name, "rb"))

            #Get end form fields from reader
            end_fields = end_form.getFields(tree=None, retval=None, fileobj=None)
            end_field_names = list(end_fields.keys())

            #Populate end field values with name and date, position depending on options
            end_field_values = [""] * 4
            index = 2 if self.new else 0
            current_date = date.today()
            end_field_values[index:index+1] = [self.client.first_name + " " + self.client.last_name, current_date.strftime("%d/%m/%Y")]

            #Zip end field values and names into dict
            end_field_dict = dict(zip(end_field_names, end_field_values))

            pdf_pages.append(pypdftk.fill_form(end_form_template_name, end_field_dict))
        pypdftk.concat(pdf_pages, os.path.join(THIS_FOLDER, "static/print.pdf"))
